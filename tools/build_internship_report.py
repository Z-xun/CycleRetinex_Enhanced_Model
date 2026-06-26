from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("2026年毕业实习报告（填写版-Python全栈）.docx")
BLANK_COMPANY = "        "


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "666666")
        borders.append(tag)
    tbl_pr.append(borders)


def set_font(run, size=12, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", size=12, bold=False, align=None, first_indent=True, spacing=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(spacing)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, name="黑体")
    return p


def add_info_line(doc, label, value=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}：{value}")
    set_font(r, size=12)
    return p


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)
    add_page_number(sec)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("实  习  报  告")
    set_font(r, size=28, bold=True, name="黑体")
    for _ in range(4):
        doc.add_paragraph()
    add_info_line(doc, "学生姓名", "")
    add_info_line(doc, "学    号", "")
    add_info_line(doc, "专业班级", "计算机科学与技术220  班")
    add_info_line(doc, "实习类别", "□认识实习  □生产实习  ☑毕业实习")
    add_info_line(doc, "实习方式", "□集中实习  ☑分散实习")
    add_info_line(doc, "实习单位", BLANK_COMPANY)
    add_info_line(doc, "指导教师", "")
    add_info_line(doc, "实习时间", "第 1 周到第 2 周，共 2 周  3月2日—3月13日")
    add_info_line(doc, "所属学院", "信息科学与工程学院")
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年3月")
    set_font(r, size=14)
    doc.add_page_break()


def add_notice(doc):
    add_heading(doc, "说  明", 1)
    items = [
        "每个学生都必须认真撰写《实习报告》。通过撰写实习报告，系统地回顾和总结实习的全过程，将实践性教学的感性认知升华到一定的理论高度，从而提高实习教学效果。",
        "实习总结要求条理清晰，内容详尽，数据准确，应符合实习大纲和实习指导书的要求。各专业对实习总结应提出一定的内容及字数要求。",
        "实习报告的质量反映了实习的质量，它是实习成绩评定的主要依据之一。实习报告需经指导教师审阅签字并给出成绩。不交实习报告者不得参加实习成绩评定。",
        "实习完成后将实习报告交学院教学办公室保存。",
    ]
    for i, text in enumerate(items, 1):
        add_para(doc, f"{i}、{text}", first_indent=False)
    doc.add_page_break()


def add_grade_table(doc):
    add_heading(doc, "实习成绩评定表", 1)
    table = doc.add_table(rows=5, cols=6)
    table.autofit = False
    set_borders(table)
    widths = [2.1, 3.0, 2.1, 3.0, 2.1, 3.6]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row = table.rows[0].cells
    vals = ["姓  名", "", "学  号", "", "成  绩", ""]
    for cell, val in zip(row, vals):
        cell.text = val
    row = table.rows[1].cells
    vals = ["专业班级", "计科220  班", "起始时间", "3月2日—3月13日", "实习单位", BLANK_COMPANY]
    for cell, val in zip(row, vals):
        cell.text = val
    table.cell(2, 0).text = "指导教师评语"
    table.cell(2, 1).merge(table.cell(4, 5))
    table.cell(2, 1).text = "\n\n\n\n\n指导教师：\n\n年    月    日"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(p.text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, size=11)
    add_para(doc, "注：成绩采用五级分制：优、良、中、及格、不及格", size=10, first_indent=False)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "目  录", 1)
    entries = [
        ("一、实习计划", "1"),
        ("二、实习目的", "2"),
        ("三、实习内容", "3"),
        ("四、实习日志", "4"),
        ("五、实习总结", "10"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2.5)
        p.paragraph_format.right_indent = Cm(2.5)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{title}{'.' * (42 - len(title))}{page}")
        set_font(r, size=12)
    doc.add_page_break()


def add_plan(doc):
    add_heading(doc, "一、实习计划", 1)
    rows = [
        ("3月2日", "了解公司业务、开发规范和项目背景，配置 Python 全栈开发环境。", "熟悉组织结构、代码仓库、需求文档和协作流程。"),
        ("3月3日", "学习 Django、FastAPI、Vue3 与 MySQL 在项目中的使用方式。", "完成本地服务启动，梳理前后端接口调用关系。"),
        ("3月4日", "参与企业内部工单与客户信息管理系统需求分析。", "整理客户、工单、统计看板等核心模块的字段和流程。"),
        ("3月5日", "完成用户登录、权限校验和基础数据接口的开发练习。", "掌握 JWT、RBAC 权限和接口文档编写方法。"),
        ("3月6日", "参与前端页面开发，完成列表、表单和查询组件联调。", "熟悉 Vue3、Element Plus、Axios 与状态管理。"),
        ("3月9日", "学习数据库建模、索引设计、数据迁移和测试数据构造。", "优化查询语句，完成部分业务表的初始化脚本。"),
        ("3月10日", "参与工单流转、客户跟进记录和消息提醒功能开发。", "理解中小型公司内部协同系统的典型业务场景。"),
        ("3月11日", "进行接口测试、异常处理、日志记录和缺陷修复。", "使用 Pytest、Postman 和 Git 分支协作提交代码。"),
        ("3月12日", "学习 Docker 部署、Nginx 反向代理和简单上线流程。", "完成测试环境部署演练并记录部署注意事项。"),
        ("3月13日", "整理实习资料、进行项目复盘和成果汇报。", "总结技术收获、问题不足和后续学习方向。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(["日期", "上午主要安排", "下午主要安排"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "D9EAF7")
    for date, am, pm in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, [date, am, pm]):
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, [2.2, 6.1, 6.1][i])
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, size=10.5, bold=(row == table.rows[0]))
    doc.add_page_break()


def add_purpose(doc):
    add_heading(doc, "二、实习目的", 1)
    paragraphs = [
        "通过毕业实习，培养学生树立理论联系实际的工作作风，以及在工作现场中将科学的理论知识加以验证、深化、巩固和充实的能力。实习过程能够帮助学生进行调查、研究、分析和解决实际问题，使课堂中学习到的计算机专业知识在真实项目环境中得到检验。",
        "1）行业认识与实践机会。计算机专业是一个持续发展和快速迭代的领域，通过毕业实习，学生可以更好地了解软件开发行业的岗位分工、项目流程和技术趋势。实习提供了实践机会，让学生将数据结构、数据库、软件工程、网络技术和 Web 开发等理论知识应用到实际项目中。",
        "2）技能提升与培养。实习能够帮助学生提升编程、系统设计、接口调试、数据库管理、团队协作等综合能力。在实习过程中，学生有机会学习并掌握新的工具和技术，这对未来进入软件开发岗位具有重要意义。",
        "3）职业规划与导向。通过实习，学生可以更清晰地了解自己的职业兴趣和能力短板，从而更好地规划未来职业发展道路。实习也为学生提供了与项目经理、开发人员和测试人员交流的机会，可以从实际工作经验中获得职业建议。",
        "4）积累经验。实习是学生进入职场的重要途径之一，通过参与真实项目和规范化开发流程，学生可以积累工作经验，增强简历竞争力，也能提前适应企业对质量、效率和沟通的要求。",
        "5）学习与成长。实习是学习和成长的过程，通过实践和经验积累，学生可以不断提升专业能力。实习也是反思和总结的机会，学生能够从项目中的问题和不足中汲取教训，不断改进自己，提高专业素养。",
        "综上所述，学生进行毕业实习的目的是多方面的。通过毕业实习，可以拓宽知识面，增加感性认识，把所学知识条理化、系统化，获得书本以外的专业经验，并进一步激发向实践学习和探索的积极性，为今后的学习和技术工作打下基础。",
    ]
    for text in paragraphs:
        add_para(doc, text)
    doc.add_page_break()


def add_content(doc):
    add_heading(doc, "三、实习内容", 1)
    paragraphs = [
        f"本次毕业实习在{BLANK_COMPANY}进行，实习项目围绕一套“企业内部工单与客户信息管理系统”展开。该项目服务于中小型公司的日常客户跟进、问题处理和内部协同，整体规模较为适中，主要由客户资料管理、工单流转、任务分配、数据统计、权限控制和消息提醒等模块组成。系统面向公司内部业务人员、技术支持人员和管理人员使用，项目目标不是追求复杂的大型平台，而是解决公司日常工作中信息分散、处理进度不透明和数据统计不及时的问题。",
        "项目采用 Python 全栈开发思路，后端以 Django 和 Django REST Framework 为主要框架，部分高频查询接口使用 FastAPI 进行练习和验证；前端采用 Vue3、Element Plus 和 Axios 完成页面交互；数据库使用 MySQL 保存业务数据，Redis 用于缓存登录状态和部分消息提醒数据；版本管理使用 Git，接口测试使用 Postman 和 Pytest，部署演练使用 Docker、Nginx 和 Gunicorn。这样的技术选型符合中小型公司项目的实际情况，既能保证开发效率，也便于后期维护。",
        "实习初期，我主要学习公司的开发规范和项目结构，配置本地开发环境，熟悉前后端分离项目的运行方式。在导师指导下，我阅读了需求说明和接口文档，了解客户信息、工单状态、处理记录和统计看板之间的业务关系。通过梳理数据表字段和页面流程，我认识到真实项目不仅需要实现功能，还要考虑权限边界、数据一致性、异常提示和操作记录等细节。",
        "在后端开发方面，我参与了用户登录、客户列表查询、工单新增与编辑、处理记录保存等接口的开发练习。通过 Django ORM 完成模型设计和数据查询，使用序列化器完成数据校验，并对接口返回格式进行统一封装。在开发过程中，我学习了 JWT 身份认证、基于角色的权限控制、分页查询、条件筛选和异常处理等内容，也认识到接口设计需要兼顾前端调用便利性和后端数据安全性。",
        "在前端开发方面，我根据已有页面风格完成了客户列表、工单表单和统计卡片等页面组件的修改。通过 Vue3 组件化开发，我学习了页面状态管理、表单校验、弹窗交互和接口联调的方法。前端开发让我意识到用户体验对业务系统同样重要，例如查询条件是否清晰、表单提示是否准确、提交失败后是否能给出明确原因，都会影响实际使用效率。",
        "在数据库与测试方面，我学习了 MySQL 表结构设计、索引使用和数据迁移脚本编写。针对客户名称、工单状态和创建时间等常用查询条件，我尝试分析查询语句并添加合适索引。随后使用 Pytest 编写简单接口测试，覆盖登录、客户查询、工单创建和权限校验等场景。通过测试和缺陷修复，我逐渐理解到稳定性和可维护性是企业项目非常重要的质量要求。",
        "实习后期，我参与了 Docker 部署演练和项目复盘。通过编写环境变量配置、启动后端服务、配置 Nginx 反向代理和检查日志，我对项目从本地开发到测试环境运行的过程有了更完整的认识。虽然实习时间较短，参与的代码量有限，但整个过程让我比较系统地体验了 Python 全栈项目从需求分析、编码实现、联调测试到部署总结的基本流程。",
    ]
    for text in paragraphs:
        add_para(doc, text)
    doc.add_page_break()


LOGS = [
    ("2026年3月2日 星期一", "了解实习单位的业务方向、团队分工和项目背景，完成 Python、Node.js、MySQL、Git 等开发环境安装，熟悉项目代码仓库和开发规范。", "第一天主要以熟悉环境和了解项目为主。我认识到企业项目对环境一致性和开发规范要求较高，不能只关注代码能否运行，还要关注分支命名、提交说明、配置文件管理和文档记录。通过搭建本地环境，我初步掌握了项目启动流程，也为后续开发打下基础。"),
    ("2026年3月3日 星期二", "学习 Django、Django REST Framework、Vue3、Element Plus 在项目中的使用方式，阅读接口文档和已有模块代码，梳理前后端调用关系。", "通过阅读代码，我对 Python 全栈项目的分层结构有了更清晰的认识。后端负责模型、接口、权限和业务规则，前端负责展示、交互和数据提交。以前学习框架时更多关注单个知识点，这次在完整项目中看到它们如何配合，理解更加具体。"),
    ("2026年3月4日 星期三", "参与企业内部工单与客户信息管理系统的需求分析，整理客户资料、工单状态、处理记录、统计看板等核心模块的业务流程。", "需求分析让我认识到写代码之前必须先理解业务。客户信息和工单处理看似简单，但涉及录入、分配、跟进、关闭和统计等多个环节。如果字段设计不清晰，后续接口和页面都会反复修改。因此我学习了用流程图和字段表记录需求的方法。"),
    ("2026年3月5日 星期四", "练习用户登录、JWT 身份认证和角色权限校验，完成部分基础接口的阅读和调试，学习统一返回格式和异常处理方式。", "今天的学习重点是后端接口安全。通过调试登录接口，我理解了 Token 的生成、携带和校验过程，也认识到不同角色应看到不同数据。权限控制如果处理不好，可能造成业务数据越权访问，因此后端开发必须保持严谨。"),
    ("2026年3月6日 星期五", "参与客户列表和工单表单页面开发，使用 Vue3、Element Plus、Axios 完成列表查询、表单校验、弹窗提交和接口联调。", "前端联调过程中，我遇到过字段命名不一致、日期格式不匹配和错误提示不明确等问题。通过与后端接口对照，我学会了先确认数据结构，再定位问题来源。页面开发不仅要让功能可用，还要让操作流程清楚、反馈及时。"),
    ("2026年3月9日 星期一", "学习 MySQL 数据库建模、索引设计和迁移脚本编写，构造测试数据，分析客户名称、工单状态、创建时间等查询条件。", "数据库设计直接影响系统后期维护。今天我认识到字段类型、默认值、索引和关联关系都需要结合业务场景考虑。通过分析查询语句，我理解了索引优化的基本思路，也知道不能盲目增加索引，要兼顾写入成本和查询频率。"),
    ("2026年3月10日 星期二", "参与工单流转、客户跟进记录和消息提醒功能开发，学习状态机思想，梳理待处理、处理中、已完成等状态变化规则。", "工单流转模块让我体会到业务规则的重要性。状态变化必须有明确条件和操作记录，否则后续很难追踪责任。通过编写部分逻辑和阅读导师的实现，我学习了如何把业务流程转化为代码中的条件判断、数据库记录和接口返回。"),
    ("2026年3月11日 星期三", "进行接口测试和缺陷修复，使用 Postman 验证接口参数，使用 Pytest 编写简单测试用例，检查异常处理和日志记录。", "测试过程让我发现很多问题并不是语法错误，而是边界条件考虑不足，例如空参数、重复提交和权限不足等情况。通过补充测试，我认识到测试能够帮助开发者提前发现风险，也能提高后续修改代码时的信心。"),
    ("2026年3月12日 星期四", "学习 Docker 部署、Gunicorn 启动后端服务、Nginx 反向代理和环境变量配置，完成测试环境部署演练并查看运行日志。", "部署演练让我第一次较完整地理解项目上线前的准备工作。程序在本地能运行并不代表在服务器上就能稳定运行，还需要考虑端口、配置、依赖、日志和静态资源等问题。运维知识能帮助开发者更好地定位线上问题。"),
    ("2026年3月13日 星期五", "整理实习资料，完成项目复盘和成果汇报，总结 Python 全栈开发流程、个人收获和后续学习计划。", "最后一天主要进行总结。我认识到自己在框架原理、数据库优化和工程化部署方面还有提升空间，但通过两周实习，我已经对中小型企业 Python 全栈项目有了整体认识。今后需要继续加强代码规范、测试能力和业务理解能力。"),
]


def add_logs(doc):
    add_heading(doc, "四、实习日志", 1)
    for date, task, feeling in LOGS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(date)
        set_font(r, size=12, bold=True, name="黑体")
        table = doc.add_table(rows=2, cols=2)
        set_borders(table)
        table.cell(0, 0).text = "实习任务（内容）"
        table.cell(0, 1).text = task
        table.cell(1, 0).text = "完成情况及心得"
        table.cell(1, 1).text = feeling
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                set_cell_width(cell, 3.0 if i == 0 else 11.5)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if i == 0:
                    set_cell_shading(cell, "EAF3F8")
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.25
                    para.paragraph_format.space_after = Pt(2)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        set_font(run, size=10.5, bold=(i == 0))
    doc.add_page_break()


def add_summary(doc):
    add_heading(doc, "五、实习总结", 1)
    paragraphs = [
        f"两周的毕业实习时间虽然不长，但我在{BLANK_COMPANY}参与 Python 全栈项目实践的过程中，对软件开发岗位和企业项目流程有了更真实的认识。本次实习的核心项目是企业内部工单与客户信息管理系统，项目规模符合中小型公司的实际需求，功能重点集中在客户资料维护、工单处理、内部协作和数据统计等方面。通过参与项目学习和部分模块开发，我把课堂上学到的编程、数据库、网络和软件工程知识与实际业务结合起来，收获比较明显。",
        "在技术方面，我进一步熟悉了 Python Web 开发的基本流程。以前学习 Django 或 Vue 时，多数是在单独的练习项目中完成某个功能，对真实项目中的目录结构、配置管理、权限控制和接口规范理解不够深入。实习中我通过阅读代码、调试接口和完成小任务，逐渐理解了后端模型、序列化器、视图函数、权限校验和前端组件之间的关系，也认识到一个稳定系统需要前后端、数据库、测试和部署共同配合。",
        "在后端学习中，我重点掌握了 Django REST Framework 的接口开发方式，了解了 JWT 登录认证、RBAC 权限控制、分页查询、条件筛选和统一异常处理。尤其是在工单流转功能中，我体会到业务规则对代码实现的影响。不同状态之间能否流转、谁有权限处理、是否需要保留操作记录，这些问题都需要在开发前想清楚。通过这部分实践，我认识到后端开发不仅是写接口，更是把业务流程严谨地转化为可维护的程序。",
        "在前端学习中，我对 Vue3 组件化开发和接口联调有了更直观的认识。客户列表、工单表单和统计卡片等页面功能看起来并不复杂，但要做到数据展示清晰、表单校验准确、错误提示友好，需要反复调试。通过处理字段格式、日期展示、弹窗提交和接口返回信息，我认识到前端开发同样需要关注用户使用场景。一个好的业务系统应当让使用者快速找到信息、明确知道下一步操作，并在出错时获得清楚提示。",
        "数据库与测试方面的收获也很大。通过参与表结构梳理和测试数据构造，我理解了数据库设计要服务于业务查询和数据维护。索引、字段约束和关联关系并不是孤立的技术点，而是影响系统性能和数据质量的重要因素。在接口测试过程中，我学会了使用 Postman 验证参数和返回结果，也尝试用 Pytest 编写简单测试用例。测试让我认识到软件质量不是最后才检查出来的，而应贯穿开发过程。",
        "在工程化方面，我学习了 Git 分支协作、代码提交规范、Docker 部署演练、Nginx 反向代理和日志检查等内容。这些内容以前接触较少，但在企业开发中非常重要。项目从本地运行到测试环境部署，需要考虑依赖、配置、端口、静态资源和日志等多个因素。通过部署演练，我对软件交付流程有了更完整的认识，也意识到开发人员具备一定运维意识能够提高问题定位效率。",
        "除了专业技能，本次实习还让我体会到团队沟通和工作习惯的重要性。实习过程中，我需要根据任务安排及时反馈进度，遇到问题先查阅文档和日志，再向导师请教。与团队成员沟通接口字段、页面效果和业务规则时，我逐渐学会用更清楚的方式描述问题。良好的沟通能够减少误解，提高协作效率，这也是课堂学习中较难直接获得的经验。",
        "当然，我也发现了自身不足。首先，对框架底层原理理解还不够深入，遇到复杂报错时定位速度较慢；其次，数据库性能优化和复杂 SQL 编写能力仍需加强；再次，对自动化测试和持续集成的实践经验还比较有限。针对这些问题，我计划在后续学习中继续深入 Python Web 框架、数据库原理和工程化部署，同时通过更多项目练习提升独立分析和解决问题的能力。",
        "总体来说，本次毕业实习是一次非常有价值的实践经历。它让我从学生视角进一步走向工程实践视角，理解了中小型公司软件项目的开发流程和岗位要求。通过 Python 全栈项目的学习，我不仅提升了技术能力，也增强了职业规划意识。今后我会继续保持主动学习和认真总结的习惯，把实习中获得的经验转化为后续毕业设计、求职和工作的基础。",
    ]
    for text in paragraphs:
        add_para(doc, text)


def main():
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_notice(doc)
    add_grade_table(doc)
    add_toc(doc)
    add_plan(doc)
    add_purpose(doc)
    add_content(doc)
    add_logs(doc)
    add_summary(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
